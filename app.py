import streamlit as st
import json
import csv
import io
import time
from pathlib import Path

from parser import parse_brief

st.set_page_config(
    page_title="BriefBot — AI Brief Parser",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

    :root {
        --bg-primary: #0a0a0f;
        --bg-secondary: #12121a;
        --bg-card: #16161f;
        --bg-elevated: #1c1c28;
        --border-subtle: rgba(255,255,255,0.06);
        --border-medium: rgba(255,255,255,0.1);
        --text-primary: #f0f0f5;
        --text-secondary: #a0a0b0;
        --text-dim: #6b6b80;
        --accent: #6366f1;
        --accent-glow: rgba(99,102,241,0.15);
        --green: #34d399;
        --green-dim: rgba(52,211,153,0.12);
        --yellow: #fbbf24;
        --yellow-dim: rgba(251,191,36,0.12);
        --red: #f87171;
        --red-dim: rgba(248,113,113,0.12);
    }

    .stApp {
        font-family: 'Inter', -apple-system, sans-serif;
        background-color: var(--bg-primary) !important;
        color: var(--text-primary);
    }

    /* Override Streamlit theme */
    [data-testid="stAppViewContainer"] {
        background-color: var(--bg-primary) !important;
    }
    [data-testid="stHeader"] {
        background: var(--bg-primary) !important;
        border-bottom: 1px solid var(--border-subtle);
    }
    [data-testid="stSidebar"] { display: none; }

    /* Override form/input backgrounds */
    .stTextArea textarea {
        background-color: var(--bg-secondary) !important;
        color: var(--text-primary) !important;
        border: 1px solid var(--border-medium) !important;
        border-radius: 12px !important;
        font-family: 'Inter', sans-serif !important;
        font-size: 0.88rem !important;
        line-height: 1.6 !important;
    }
    .stTextArea textarea:focus {
        border-color: var(--accent) !important;
        box-shadow: 0 0 0 3px var(--accent-glow) !important;
    }
    .stTextArea textarea::placeholder { color: var(--text-dim) !important; }

    /* File uploader */
    [data-testid="stFileUploader"] {
        background: var(--bg-secondary) !important;
        border: 1px dashed var(--border-medium) !important;
        border-radius: 12px !important;
    }
    [data-testid="stFileUploader"] label { color: var(--text-secondary) !important; }
    [data-testid="stFileUploader"] span { color: var(--text-dim) !important; }
    [data-testid="stFileUploader"] small { color: var(--text-dim) !important; }

    /* Buttons */
    button[kind="secondary"] {
        background: var(--bg-elevated) !important;
        color: var(--text-secondary) !important;
        border: 1px solid var(--border-medium) !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        transition: all 0.15s ease !important;
    }
    button[kind="secondary"]:hover {
        background: var(--bg-card) !important;
        border-color: var(--accent) !important;
        color: var(--text-primary) !important;
    }
    button[kind="secondary"][aria-pressed="true"],
    button[kind="secondary"]:active {
        background: var(--accent-glow) !important;
        border-color: var(--accent) !important;
        color: var(--accent) !important;
    }
    button[kind="primary"] {
        background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        padding: 0.75rem 1.5rem !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 4px 16px rgba(99,102,241,0.3) !important;
    }
    button[kind="primary"]:hover {
        box-shadow: 0 6px 24px rgba(99,102,241,0.4) !important;
        transform: translateY(-1px) !important;
    }

    /* Download buttons */
    [data-testid="stDownloadButton"] button {
        background: var(--bg-elevated) !important;
        color: var(--text-secondary) !important;
        border: 1px solid var(--border-medium) !important;
        border-radius: 8px !important;
    }

    /* Hero section */
    .hero {
        background: var(--bg-secondary);
        border: 1px solid var(--border-subtle);
        border-radius: 16px;
        padding: 2rem 2.5rem;
        margin-bottom: 1.5rem;
        position: relative;
        overflow: hidden;
    }
    .hero::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0; height: 2px;
        background: linear-gradient(90deg, #6366f1, #8b5cf6, #6366f1);
    }
    .hero h1 {
        font-size: 1.8rem;
        font-weight: 800;
        margin: 0;
        letter-spacing: -0.5px;
        color: var(--text-primary);
    }
    .hero p {
        font-size: 0.95rem;
        color: var(--text-secondary);
        margin: 0.4rem 0 0 0;
        font-weight: 300;
    }

    /* Stats bar */
    .stats-bar {
        display: flex;
        gap: 0.75rem;
        margin: 1rem 0;
    }
    .stat-item {
        background: var(--bg-card);
        border: 1px solid var(--border-subtle);
        border-radius: 12px;
        padding: 0.85rem 1rem;
        flex: 1;
        text-align: center;
    }
    .stat-number {
        font-size: 1.5rem;
        font-weight: 700;
        background: linear-gradient(135deg, #6366f1, #a78bfa);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        line-height: 1.1;
    }
    .stat-label {
        font-size: 0.65rem;
        text-transform: uppercase;
        letter-spacing: 1.2px;
        color: var(--text-dim);
        margin-top: 0.25rem;
        font-weight: 500;
    }

    /* Input section */
    .input-header {
        font-size: 0.7rem;
        text-transform: uppercase;
        letter-spacing: 1.5px;
        color: var(--text-dim);
        font-weight: 600;
        margin-bottom: 0.75rem;
    }

    /* Field cards */
    .field-card {
        background: var(--bg-card);
        border-radius: 10px;
        padding: 0.85rem 1.1rem;
        margin: 0.4rem 0;
        border: 1px solid var(--border-subtle);
        border-left: 3px solid var(--border-medium);
        transition: all 0.15s ease;
    }
    .field-card:hover {
        border-color: var(--border-medium);
        background: var(--bg-elevated);
    }
    .field-card-high { border-left-color: var(--green); }
    .field-card-medium { border-left-color: var(--yellow); }
    .field-card-low { border-left-color: var(--red); }

    .field-label {
        font-size: 0.65rem;
        text-transform: uppercase;
        letter-spacing: 1px;
        color: var(--text-dim);
        font-weight: 600;
    }
    .field-value {
        font-size: 0.88rem;
        color: var(--text-primary);
        margin-top: 0.25rem;
        line-height: 1.5;
    }
    .field-badge {
        float: right;
        font-size: 0.6rem;
        padding: 2px 8px;
        border-radius: 10px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .badge-high { background: var(--green-dim); color: var(--green); }
    .badge-medium { background: var(--yellow-dim); color: var(--yellow); }
    .badge-low { background: var(--red-dim); color: var(--red); }

    /* Ambiguity section */
    .ambiguity-section {
        background: var(--bg-card);
        border: 1px solid rgba(251,191,36,0.2);
        border-radius: 12px;
        padding: 1.1rem 1.25rem;
        margin: 1rem 0;
    }
    .ambiguity-header {
        font-size: 0.85rem;
        font-weight: 700;
        color: var(--yellow);
        margin-bottom: 0.6rem;
    }
    .ambiguity-item {
        font-size: 0.82rem;
        color: var(--text-secondary);
        padding: 0.35rem 0;
        padding-left: 0.85rem;
        border-left: 2px solid rgba(251,191,36,0.4);
        margin: 0.35rem 0;
        line-height: 1.45;
    }

    /* Deliverables */
    .deliverable-item {
        display: flex;
        align-items: center;
        padding: 0.55rem 0.9rem;
        background: var(--bg-secondary);
        border-radius: 8px;
        margin: 0.35rem 0;
        font-size: 0.83rem;
        border: 1px solid var(--border-subtle);
        color: var(--text-primary);
    }
    .deliverable-bullet {
        width: 5px;
        height: 5px;
        background: var(--accent);
        border-radius: 50%;
        margin-right: 0.7rem;
        flex-shrink: 0;
    }
    .deliverable-format {
        margin-left: auto;
        font-size: 0.68rem;
        background: var(--accent-glow);
        color: #a78bfa;
        padding: 2px 8px;
        border-radius: 6px;
        font-weight: 500;
        white-space: nowrap;
    }

    /* Project header */
    .project-title {
        font-size: 1.35rem;
        font-weight: 700;
        color: var(--text-primary);
        margin: 0;
        line-height: 1.2;
    }
    .project-meta {
        font-size: 0.78rem;
        color: var(--text-secondary);
        margin-top: 0.4rem;
        display: flex;
        gap: 0.75rem;
        flex-wrap: wrap;
    }
    .meta-tag {
        background: var(--bg-elevated);
        border: 1px solid var(--border-subtle);
        padding: 3px 10px;
        border-radius: 6px;
        font-weight: 500;
        color: var(--text-secondary);
    }

    /* Export section */
    .export-section {
        background: var(--bg-secondary);
        border-radius: 12px;
        padding: 1rem;
        margin-top: 1.25rem;
        border: 1px solid var(--border-subtle);
    }

    /* Section dividers */
    .section-divider {
        height: 1px;
        background: linear-gradient(to right, transparent, var(--border-medium), transparent);
        margin: 1.25rem 0;
    }

    /* Column gap styling */
    [data-testid="stVerticalBlock"] { gap: 0.5rem !important; }

    /* Spinner/loading */
    .stSpinner > div { border-top-color: var(--accent) !important; }

    /* Hide streamlit branding */
    #MainMenu, footer { display: none !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="hero">
    <h1>BriefBot</h1>
    <p>Paste a messy creative brief. Get structured requirements in seconds — with smart follow-up questions your team would forget to ask.</p>
</div>
""", unsafe_allow_html=True)

EXAMPLES_DIR = Path(__file__).parent / "examples"


def load_example(filename):
    return (EXAMPLES_DIR / filename).read_text()


def confidence_badge_html(level):
    badge_class = f"badge-{level}" if level in ("high", "medium", "low") else ""
    return f'<span class="field-badge {badge_class}">{level}</span>'


def format_jira_template(result):
    lines = []
    lines.append(f"h2. {result.get('project_name', 'Untitled Project')}")
    lines.append(f"*Brand:* {result.get('brand', 'N/A')}")
    lines.append(f"*Deadline:* {result.get('deadline', 'Not specified')}")
    lines.append(f"*Target Audience:* {result.get('target_audience', 'N/A')}")
    lines.append("")
    lines.append("h3. Key Message")
    lines.append(result.get("key_message", "N/A"))
    lines.append(f"*CTA:* {result.get('cta', 'N/A')}")
    lines.append("")
    lines.append("h3. Deliverables")
    for d in result.get("deliverables", []):
        desc = d.get("description", "")
        fmt = d.get("format", "")
        qty = d.get("quantity", "")
        detail = f" ({fmt})" if fmt else ""
        detail += f" x{qty}" if qty else ""
        lines.append(f"* {desc}{detail}")
    lines.append("")
    lines.append("h3. Dimensions")
    for dim in result.get("dimensions", []):
        lines.append(f"* {dim}")
    lines.append("")
    lines.append("h3. Brand Guidelines")
    for g in result.get("brand_guidelines", []):
        lines.append(f"* {g}")
    lines.append("")
    lines.append("h3. Channels")
    lines.append(", ".join(result.get("channels", [])))
    lines.append("")
    if result.get("ambiguities"):
        lines.append("h3. {color:red}Open Questions{color}")
        for a in result["ambiguities"]:
            lines.append(f"* {a}")
    return "\n".join(lines)


col_input, col_gap, col_output = st.columns([5, 0.5, 6])

with col_input:
    st.markdown('<p class="input-header">Input — Paste or Upload</p>', unsafe_allow_html=True)

    example_cols = st.columns(3)
    with example_cols[0]:
        if st.button("📧 Messy Email", use_container_width=True):
            st.session_state["brief_input"] = load_example("messy_email.txt")
            st.session_state.pop("result", None)
            st.rerun()
    with example_cols[1]:
        if st.button("📄 Word Doc", use_container_width=True):
            from docx import Document
            doc = Document(Path(__file__).parent / "examples" / "word_brief.docx")
            text = "\n".join(p.text for p in doc.paragraphs)
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(" | ".join(cell.text for cell in row.cells))
            st.session_state["brief_input"] = text + "\n\n" + "\n".join(tables)
            st.session_state.pop("result", None)
            st.rerun()
    with example_cols[2]:
        if st.button("💬 Slack Thread", use_container_width=True):
            st.session_state["brief_input"] = load_example("slack_thread.txt")
            st.session_state.pop("result", None)
            st.rerun()

    brief_text = st.text_area(
        "Brief content:",
        height=460,
        key="brief_input",
        label_visibility="collapsed",
        placeholder="Paste your creative brief here — emails, Slack messages, bullet points, anything..."
    )

    uploaded_file = st.file_uploader(
        "Or upload a file:",
        type=["txt", "pdf", "docx", "xlsx", "xls", "csv"],
        label_visibility="collapsed"
    )

    if uploaded_file:
        if uploaded_file.type == "text/plain":
            brief_text = uploaded_file.read().decode("utf-8")
        elif uploaded_file.type == "text/csv" or uploaded_file.name.endswith(".csv"):
            try:
                import pandas as pd
                df = pd.read_csv(uploaded_file)
                brief_text = df.to_string(index=False)
            except Exception as e:
                st.error(f"Error reading CSV: {e}")
        elif uploaded_file.type == "application/pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(uploaded_file)
                brief_text = "\n".join(page.extract_text() for page in reader.pages)
            except Exception as e:
                st.error(f"Error reading PDF: {e}")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                from docx import Document
                doc = Document(uploaded_file)
                brief_text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                st.error(f"Error reading DOCX: {e}")
        elif uploaded_file.type in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel") or uploaded_file.name.endswith((".xlsx", ".xls")):
            try:
                import pandas as pd
                df = pd.read_excel(uploaded_file)
                brief_text = df.to_string(index=False)
            except Exception as e:
                st.error(f"Error reading spreadsheet: {e}")

    parse_clicked = st.button(
        "⚡ Parse Brief",
        type="primary",
        use_container_width=True
    )

with col_output:
    if parse_clicked and brief_text.strip():
        with st.spinner(""):
            st.markdown("""
                <div style="text-align: center; padding: 3rem 0; color: var(--text-dim);">
                    <div style="font-size: 1.5rem; margin-bottom: 0.5rem; opacity: 0.7;">&#9889;</div>
                    <div style="font-size: 0.85rem; font-weight: 500; color: var(--text-secondary);">Analyzing brief with Claude...</div>
                    <div style="font-size: 0.72rem; margin-top: 0.3rem; color: var(--text-dim);">Extracting requirements, flagging ambiguities</div>
                </div>
            """, unsafe_allow_html=True)
            start_time = time.time()
            result = parse_brief(brief_text)
            elapsed = time.time() - start_time
            st.session_state["result"] = result
            st.session_state["elapsed"] = elapsed
            st.rerun()

    if "result" in st.session_state:
        result = st.session_state["result"]
        elapsed = st.session_state.get("elapsed", 0)
        confidence = result.get("confidence", {})

        if "error" in result:
            st.error(result["error"])
        else:
            # Stats bar
            high_count = sum(1 for v in confidence.values() if v == "high")
            medium_count = sum(1 for v in confidence.values() if v == "medium")
            ambiguity_count = len(result.get("ambiguities", []))
            deliverable_count = len(result.get("deliverables", []))

            st.markdown(f"""
            <div class="stats-bar">
                <div class="stat-item">
                    <div class="stat-number">{elapsed:.1f}s</div>
                    <div class="stat-label">Parse Time</div>
                </div>
                <div class="stat-item">
                    <div class="stat-number">{deliverable_count}</div>
                    <div class="stat-label">Deliverables</div>
                </div>
                <div class="stat-item">
                    <div class="stat-number">{high_count}/{high_count + medium_count}</div>
                    <div class="stat-label">High Confidence</div>
                </div>
                <div class="stat-item">
                    <div class="stat-number">{ambiguity_count}</div>
                    <div class="stat-label">Questions Found</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Agent review metadata
            agent_meta = result.get("_agent", {})
            if agent_meta.get("passes", 1) > 1:
                review_status = agent_meta.get("review", "unknown")
                corrections = agent_meta.get("corrections_applied", 0)
                downgrades = agent_meta.get("confidence_downgrades", 0)
                added = agent_meta.get("ambiguities_added", 0)
                status_color = "#10b981" if review_status == "pass" else "#f59e0b" if review_status == "minor_corrections" else "#ef4444"
                status_label = "Verified" if review_status == "pass" else "Corrected" if review_status == "minor_corrections" else "Significantly revised"
                corrections_text = ""
                if corrections > 0 or downgrades > 0 or added > 0:
                    parts = []
                    if corrections > 0:
                        parts.append(f"{corrections} field{'s' if corrections != 1 else ''} corrected")
                    if downgrades > 0:
                        parts.append(f"{downgrades} confidence{'s' if downgrades != 1 else ''} adjusted")
                    if added > 0:
                        parts.append(f"{added} question{'s' if added != 1 else ''} added")
                    corrections_text = " — " + ", ".join(parts)
                st.markdown(f"""
                <div style="background: rgba(99,102,241,0.06); border: 1px solid rgba(99,102,241,0.2); border-radius: 10px; padding: 0.55rem 1rem; margin-bottom: 0.75rem; display: flex; align-items: center; gap: 0.5rem;">
                    <span style="color: {status_color}; font-weight: 600; font-size: 0.8rem;">Agent Review: {status_label}</span>
                    <span style="color: var(--text-dim); font-size: 0.72rem;">(2-pass extraction + self-review{corrections_text})</span>
                </div>
                """, unsafe_allow_html=True)

            # Project header
            deadline = result.get('deadline', 'Not specified')
            brand = result.get('brand', 'Unknown')
            st.markdown(f"""
            <div style="margin-bottom: 1.25rem;">
                <p class="project-title">{result.get('project_name', 'Untitled')}</p>
                <div class="project-meta">
                    <span class="meta-tag">🏢 {brand}</span>
                    <span class="meta-tag">📅 {deadline}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Core fields
            fields = [
                ("Target Audience", "target_audience", "👥"),
                ("Key Message", "key_message", "💬"),
                ("Call to Action", "cta", "🎯"),
                ("Asset Types", "asset_type", "🎨"),
                ("Channels", "channels", "📡"),
                ("Dimensions", "dimensions", "📐"),
                ("Brand Guidelines", "brand_guidelines", "🎨"),
                ("Stakeholders", "stakeholders", "👤"),
            ]

            for label, key, icon in fields:
                conf = confidence.get(key, "unknown")
                conf_class = f"field-card-{conf}" if conf in ("high", "medium", "low") else ""
                value = result.get(key, "N/A")
                if isinstance(value, list):
                    value_str = " · ".join(str(v) for v in value)
                else:
                    value_str = str(value)
                st.markdown(
                    f'<div class="field-card {conf_class}">'
                    f'{confidence_badge_html(conf)}'
                    f'<div class="field-label">{icon} {label}</div>'
                    f'<div class="field-value">{value_str}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            # Deliverables
            if result.get("deliverables"):
                st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
                st.markdown('<p class="input-header">📦 Deliverables</p>', unsafe_allow_html=True)
                for d in result["deliverables"]:
                    desc = d.get("description", "")
                    fmt = d.get("format", "")
                    fmt_html = f'<span class="deliverable-format">{fmt}</span>' if fmt else ""
                    st.markdown(
                        f'<div class="deliverable-item">'
                        f'<div class="deliverable-bullet"></div>'
                        f'{desc}{fmt_html}'
                        f'</div>',
                        unsafe_allow_html=True
                    )

            # Ambiguities
            if result.get("ambiguities"):
                st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
                amb_html = "".join(
                    f'<div class="ambiguity-item">{amb}</div>'
                    for amb in result["ambiguities"]
                )
                st.markdown(
                    f'<div class="ambiguity-section">'
                    f'<div class="ambiguity-header">⚠️ {ambiguity_count} Questions to Clarify Before Starting</div>'
                    f'{amb_html}'
                    f'</div>',
                    unsafe_allow_html=True
                )

            # Export
            st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
            st.markdown('<div class="export-section">', unsafe_allow_html=True)
            st.markdown('<p class="input-header">Export</p>', unsafe_allow_html=True)
            export_cols = st.columns(3)

            with export_cols[0]:
                json_str = json.dumps(result, indent=2)
                st.download_button(
                    "📥 JSON",
                    data=json_str,
                    file_name="parsed_brief.json",
                    mime="application/json",
                    use_container_width=True
                )

            with export_cols[1]:
                output = io.StringIO()
                flat = {k: json.dumps(v) if isinstance(v, (list, dict)) else v for k, v in result.items()}
                writer = csv.DictWriter(output, fieldnames=flat.keys())
                writer.writeheader()
                writer.writerow(flat)
                st.download_button(
                    "📥 CSV",
                    data=output.getvalue(),
                    file_name="parsed_brief.csv",
                    mime="text/csv",
                    use_container_width=True
                )

            with export_cols[2]:
                jira_text = format_jira_template(result)
                st.download_button(
                    "📋 Jira",
                    data=jira_text,
                    file_name="brief_jira_template.txt",
                    mime="text/plain",
                    use_container_width=True
                )

            st.markdown('</div>', unsafe_allow_html=True)

    elif parse_clicked:
        st.warning("Please paste or upload a brief first.")
    else:
        st.markdown("""
        <div style="text-align: center; padding: 4rem 2rem;">
            <div style="font-size: 2.5rem; margin-bottom: 0.75rem; opacity: 0.3;">&#9889;</div>
            <div style="font-size: 1rem; font-weight: 500; color: var(--text-dim);">No brief parsed yet</div>
            <div style="font-size: 0.82rem; margin-top: 0.4rem; color: var(--text-dim); opacity: 0.7;">
                Paste a creative brief on the left or try one of the examples.<br/>
                Results will appear here with confidence scores and smart follow-up questions.
            </div>
        </div>
        """, unsafe_allow_html=True)
